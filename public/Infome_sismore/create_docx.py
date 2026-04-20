from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set base font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

def add_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('GOBIERNO REGIONAL DE UCAYALI\nDIRECCION REGIONAL DE EDUCACION\nOFICINA DE PLANEAMIENTO Y PRESUPUESTO\n')
    run.bold = True
    run.font.size = Pt(14)
    run2 = p.add_run('"PROGRAMA DE REDUCCIÓN DE LA VULNERABILIDAD Y ATENCIÓN DE EMERGENCIAS Y DESASTRES – PREVAED N° 068"')
    run2.font.size = Pt(10)

# PAGE 1
add_header(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\nANEXO N° 03\nTÉRMINOS DE REFERENCIA PARA LA CONTRATACIÓN\nDE LOCADOR DE SERVICIO')
r.bold = True
r.font.size = Pt(12)

def add_section(doc, title, content):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    if content:
        p.add_run('\n' + content)

add_section(doc, '1. DENOMINACION DE LA CONTRATACION', 'La Contratación de Requerimiento de un locador de servicio COMO SOPORTE TÉCNICO EN MATERIA DE LA GESTIÓN DE RIESGO DE DESASTRES DEL PPR 0068 PREVAED EN LA UGEL DE CORONEL PORTILLO (DISTRITOS DE IPARIA Y MASISEA) DE LA DRE UCAYALI.')
add_section(doc, '2. SOLICITANTE.', 'PPR 0068 PREVAED - OPP')

p = doc.add_paragraph()
p.add_run('3. INFORMACION PRESUPUESTARIA').bold = True
doc.add_paragraph('Rubro 00 Recurso Ordinario', style='List Bullet')
doc.add_paragraph('Meta 003', style='List Bullet')
doc.add_paragraph('Especifica de Gastos 2.3.2.9.1.1', style='List Bullet')

add_section(doc, '4. FINALIDAD PUBLICA', 'La presente contratación tiene como finalidad contratar el servicio COMO SOPORTE TÉCNICO EN MATERIA DE LA GESTIÓN DE RIESGO DE DESASTRES DEL PPR 0068 PREVAED EN LA UGEL DE CORONEL PORTILLO (DISTRITOS DE IPARIA Y MASISEA) DE LA DRE UCAYALI.')

add_section(doc, '5. DESCRIPCION DEL SERVICIO', 'El servicio contratado, consistirá en las siguientes actividades y servicios por ejecutar por el periodo hasta 80 días calendarios.')
doc.add_paragraph('Detalle de las actividades N° 01 a desarrollar:')
doc.add_paragraph('Elaboración de la campaña comunicacional denominada: Mas educación, menos riesgos en las instituciones educativas de la región Ucayali frente al chikunguya, zika, y dengue, para un inicio del año escolar seguro.', style='List Bullet')
doc.add_paragraph('Taller de capacitación para las I.E: Rol de brigadas y uso de kit para la respuesta ante emergencias y desastres, evaluación de daños y análisis de necesidades.', style='List Bullet')

doc.add_paragraph('Detalle de las actividades N° 02 a desarrollar:')
doc.add_paragraph('Elaboración de la campaña comunicacional: Preparación y respuesta oportuna a través de simulacros escolares ante fenómenos naturales y antrópicos.', style='List Bullet')
doc.add_paragraph('Taller de capacitación a directivos y docentes en organización, ejecución y evaluación del 1°simulacro nacional multipeligro.', style='List Bullet')

doc.add_paragraph('Detalle de las actividades N° 03 a desarrollar:')
doc.add_paragraph('Verificación in situ a las I.E priorizadas para la implementación de planos y dispositivos de seguridad.', style='List Bullet')
doc.add_paragraph('Elaboración y/o actualización del Plan de GRD y actividades de contingencia frente a peligros naturales y antrópicos en la región Ucayali.', style='List Bullet')

add_section(doc, '6. PERFIL DEL LOCADOR DE SERVICIO', 'El proveedor deberá cumplir los siguientes requisitos y acreditarlos, al momento de formalizar el contrato (Orden de Servicio):')

# ADD PAGE BREAK
doc.add_page_break()

# PAGE 2
add_header(doc)

doc.add_paragraph('Bachiller en Ingeniera de Sistemas o Ingeniera de Software.', style='List Bullet')
doc.add_paragraph('Experiencia profesional no menor de un (01) año en el Sector Público y/o Privado.', style='List Bullet')
doc.add_paragraph('Contar con RNP vigente', style='List Bullet')
doc.add_paragraph('Estar en condición de Activo y Habido en la SUNAT.', style='List Bullet')
doc.add_paragraph('Contar con Registro Nacional de Proveedores (RNP).', style='List Bullet')
doc.add_paragraph('Contar con código de cuenta interbancaria (CCI) - cuenta relacionada al número de RUC.', style='List Bullet')
doc.add_paragraph('Contar con Registro Único de Contribuyentes (RUC) vigente.', style='List Bullet')

add_section(doc, '7. PLAZO DE EJECUCION', 'El servicio se realizará en un plazo no mayor a 80 días calendarios como máximo, contabilizado a partir de la notificación de la orden de servicio.')

add_section(doc, '8. ACTIVIDADES A PRESENTAR', 'Se presentará un informe al término del periodo de acuerdo al siguiente detalle:')
doc.add_paragraph('Informe de Actividades N° 01: Presentación de Informe de actividades detalladas. 20 días calendario.', style='List Bullet')
doc.add_paragraph('Informe de Actividades N° 02: Presentación de Informe de actividades detalladas. 50 días calendario.', style='List Bullet')
doc.add_paragraph('Informe de Actividades N° 03: Presentación de Informe de actividades detalladas. 80 días calendario.', style='List Bullet')

add_section(doc, '9. FORMA DE PAGO', 'El pago se realizará al término del servicio y la conformidad correspondiente, de acuerdo al siguiente detalle:')
doc.add_paragraph('El pago del Informe de las Actividades N°01 es de S/ 2,500.00 (Dos Mil Quinientos con 00/100) soles.', style='List Bullet')
doc.add_paragraph('El pago del Informe de las Actividades N°02 es de S/ 2,500.00 (Dos Mil Quinientos con 00/100) soles.', style='List Bullet')
doc.add_paragraph('El pago del Informe de las Actividades N°03 es de S/ 2,500.00 (Dos Mil Quinientos con 00/100) soles.', style='List Bullet')

add_section(doc, '10. PENALIDAD', 'En caso de retraso injustificado en la ejecución de las prestaciones objeto de la orden, se aplica automáticamente una penalidad por mora por cada día de retraso, calculado de acuerdo a la siguiente formula:\n\nPenalidad diaria = (0.10 x Monto) / (F x Plazo en días)\n\nDonde:\nF=0.25 para plazos mayores a sesenta (60) días para o\nF=0.40 para plazos menores o iguales a sesenta (60)días\nEl monto máximo de la penalidad por mora aplicable no puede exceder del diez por ciento (10%) del monto total contratado. La entidad tiene el derecho a exigir, además de la penalidad, el cumplimiento de la Obligación.')

add_section(doc, '11. CONFORMIDAD DEL SERVICIO', 'La conformidad del servicio será otorgada por el jefe y/o encargado de la Oficina de Planeamiento y Presupuesto de la sede regional.')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('\nPucallpa 09 de Abril del 2026\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Mg. Victor George Garcia López\nDIRECTOR DE LA OFICINA DE\nPLANEAMIENTO Y PRESUPUESTO DREU')
r.bold = True

# Add footer for Both pages
# In docx, footer is typically added to sections. We have two pages, but they are likely within the same section.
section = doc.sections[0]
footer = section.footer
ftr_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ftr = ftr_p.add_run('"Por una Gestión Educativa de calidad Diversificada, con Valores para el Desarrollo de la Región Ucayali"\n\nAv. Sáenz Peña N° 220 Pucallpa – Perú\nTeléfono (51)(61) 57-1433 Telefax 57-2236  E-mail: dreucayali_2@hotmail.com')
r_ftr.font.size = Pt(8)
r_ftr.bold = True

doc.save('Transcripcion_TDR.docx')
print("Document saved successfully!")
