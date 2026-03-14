import docx
import sys

def inspect_docx(filepath):
    doc = docx.Document(filepath)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        has_image = "graphic" in p._element.xml
        
        if text or has_image:
            print(f"[{i}] TEXT: {text[:100]}")
            if has_image:
                print(f"    --> HAS IMAGE")

if __name__ == '__main__':
    inspect_docx(sys.argv[1])
