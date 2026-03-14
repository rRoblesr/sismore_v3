import docx
from docx.shared import Pt
import sys

def insert_paragraph_after(paragraph, text):
    new_p_xml = docx.oxml.OxmlElement('w:p')
    paragraph._p.addnext(new_p_xml)
    new_p = docx.text.paragraph.Paragraph(new_p_xml, paragraph._parent)
    run = new_p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    return new_p

def process():
    filepath = sys.argv[1]
    outpath = sys.argv[2]
    
    descriptions = {
        "REALIZAR LA ACTUALIZACIÓN DEL PADRON WEB EN LA PLATAFORMA INFORMÁTICA DEL SISMORE EDUCACIÓN, MOSEVA": 
            "Esta actividad consistió en la actualización y sincronización del Padrón Web en múltiples plataformas del sistema educativo regional. El trabajo incluyó la carga y actualización de la información en la plataforma principal SISMORE Educación, asegurando que los datos de beneficiarios del Programa Presupuestal por Resultados (PPR) 0068 estuvieran correctamente registrados y disponibles para consulta. Asimismo, se implementó la integración del Padrón Web actualizado con el módulo MOSEVA para garantizar la coherencia de datos, y se configuró el panel de control para mostrar estadísticas actualizadas.",
        "REALIZAR LA ACTUALIZACIÓN DEL TABLERO DE CONTROL Y REPORTES DEL SISTEMA DE CONTROL Y SEGUIMIENTO DE":
            "Esta actividad involucró el mantenimiento y actualización del tablero de control y reportes del Sistema de Control y Seguimiento de Plazas NEXUS. Se realizó la carga y procesamiento de información actualizada sobre plazas docentes, incluyendo plazas ocupadas, vacantes y en proceso de asignación. Se mejoró la interfaz interactiva para facilitar la visualización del estado de plazas mediante filtros específicos y se configuraron reportes detallados que permiten a los responsables monitorear la distribución y generar estadísticas precisas sobre la situación de las plazas.",
        "REALIZAR LA ACTUALIZACIÓN DE LA MATRICULA EDUCATIVA DEL SIAGIE EN LA PLATAFORMA INFORMÁTICA DEL SISM":
            "Esta actividad se enfocó en la actualización de datos de matrícula y la sistematización de indicadores clave del programa educativo. Se realizó la importación y procesamiento de datos de matrícula educativa provenientes del SIAGIE, actualizando la base de datos de SISMORE Educación. Además, se implementó el seguimiento y registro de las cuatro actividades trazadoras del PPR 0068, lo cual incluyó la configuración de módulos de captura de datos, consolidación de la información y la generación de reportes automáticos para medir el avance del programa en tiempo real.",
        "REALIZAR LA ACTUALIZACIÓN DEL REGISTRO DE LA INFORMACIÓN DEL SANEAMIENTO FÍSICO LEGAL DE LOS LOCALES":
            "Esta actividad comprendió el registro y la actualización integral de la base de datos relacionada con el saneamiento físico legal de los locales educativos públicos en la región. Se procesó la información técnica y legal de los predios, incorporando estos datos al sistema SISMORE de manera estructurada. Asimismo, se implementó un módulo de seguimiento que permite visualizar la situación legal de cada institución, generando reportes de estado y facilitando el monitoreo continuo para promover la formalización de los terrenos escolares.",
        "REALIZAR LA ACTUALIZACIÓN DEL PADRÓN DE INSTITUCIONES EDUCATIVAS INTERCULTURAL BILINGÜE (EIB) EN LA":
            "Esta actividad se centró en la actualización y consolidación del padrón oficial de Instituciones Educativas Intercultural Bilingüe (EIB) dentro de la plataforma SISMORE Educación. Se llevó a cabo la validación y carga de la información más reciente de las escuelas categorizadas como EIB, asegurando que se reflejen los niveles de dominio lingüístico y las formas de atención pedagógica. Se desarrollaron opciones de visualización que permiten filtrar y listar estas instituciones de manera ágil, brindando información oportuna para la toma de decisiones en el ámbito de la educación intercultural."
    }

    captions = [
        "Figura 1: Captura de pantalla de la actualización del Padrón Web en la plataforma SISMORE Educación.",
        "Figura 2: Vista del módulo MOSEVA con la integración del Padrón Web actualizado.",
        "Figura 3: Interfaz del sistema NEXUS mostrando el control y seguimiento de plazas docentes actualizadas.",
        "Figura 4: Tablero de control con las estadísticas correspondientes a plazas ocupadas y vacantes.",
        "Figura 5: Vista de los reportes generados del sistema de control y seguimiento de plazas NEXUS.",
        "Figura 6: Interfaz de actualización de la matrícula educativa del SIAGIE en SISMORE Educación.",
        "Figura 7: Visualización de la sistematización de las acciones de las cuatro actividades trazadoras del PPR 0068.",
        "Figura 8: Módulo de registro de información del Saneamiento Físico Legal de los locales educativos públicos en SISMORE.",
        "Figura 9: Reporte y seguimiento de la situación legal de los predios escolares en la plataforma.",
        "Figura 10: Interfaz de actualización del padrón de Instituciones Educativas Intercultural Bilingüe (EIB) en SISMORE Educación.",
        "Figura 11: Visualización y filtrado del padrón EIB registrado en la plataforma informática."
    ]
    
    doc = docx.Document(filepath)
    
    current_img_idx = 0
    
    paragraphs = list(doc.paragraphs)
    
    for p in paragraphs:
        text = p.text.strip()
        
        # Insert descriptions after headings
        for k, v in descriptions.items():
            if text.startswith(k):
                print(f"Matched heading: {k[:30]}...")
                insert_paragraph_after(p, v)
                break
                
        # Insert captions after images
        images_in_p = len(p._element.xpath('.//pic:pic'))
        if images_in_p > 0:
            for i in range(images_in_p):
                print(f"Found image. Inserting caption: {captions[current_img_idx][:30]}...")
                p = insert_paragraph_after(p, captions[current_img_idx])
                current_img_idx += 1
                if current_img_idx >= len(captions):
                    break
            if current_img_idx >= len(captions):
                # Only warning here 
                pass
        
    doc.save(outpath)
    print(f"\nSaved modified docx to: {outpath}")
    print(f"Total images processed: {current_img_idx}")

if __name__ == '__main__':
    process()
